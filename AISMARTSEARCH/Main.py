import ast
import csv
import os

import numpy as np
import pandas as pd
from docx import Document
from dotenv import find_dotenv, load_dotenv
from langchain.document_loaders import Docx2txtLoader, PyPDFLoader
from langchain_chroma import Chroma
from langchain_openai import OpenAI, OpenAIEmbeddings
from langchain_text_splitters import (CharacterTextSplitter,
                                      MarkdownTextSplitter,
                                      RecursiveCharacterTextSplitter,
                                      TokenTextSplitter)

"""

add worddoc parse
and after taht make it so i can parse imges but thats for later lol


"""
from sklearn.metrics.pairwise import cosine_similarity


def SaveVector(VectorEmbedList):
    with open('vector.csv', 'w', newline='') as f:
        writer = csv.writer(f)
        for value in VectorEmbedList:
            writer.writerow([value])


def ReadFromFile(file_path):
    df = pd.read_csv(file_path, header=None)
    vector_from_csv = df.values.flatten().tolist()
    return vector_from_csv


class Parsers():
    def __init__(self, apikey):
        self.apikey = apikey
        self.splitter = TokenTextSplitter(
            chunk_size=400, chunk_overlap=50, length_function=len)
        # Note make so the user chose it  spliter
        # and chunk size / overlap
        self.embedingAPI = OpenAIEmbeddings(
            openai_api_key=self.apikey, model="text-embedding-3-large")

    def SetSpliter(self, spliter, chuncksize, overlap):
        match spliter:
            case "CharacterTextSplitter":
                self.splitter = CharacterTextSplitter(
                    chunk_size=chuncksize, chunk_overlap=overlap)
            case "RecursiveCharacterTextSplitter":
                self.splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chuncksize, chunk_overlap=overlap)
            case "TokenTextSplitter":
                self.splitter = TokenTextSplitter(
                    chunk_size=chuncksize, chunk_overlap=overlap)
            case "MarkdownHeaderTextSplitter":
                self.splitter = MarkdownTextSplitter(
                    chunk_size=chuncksize, chunk_overlap=overlap)
            case _:
                self.splitter = CharacterTextSplitter(
                    chunk_size=chuncksize, chunk_overlap=overlap)

    def Print(self, showList):
        for vector in showList:
            print(str(vector)[:100]+'top')
            print(str(vector)[len(showList)-100:]+'bottem')

    def load_word_document(self, file_path):
        """Extract text from a Word (.docx) file and return as LangChain documents."""
        doc = Document(file_path)
        text = "\n".join(
            [para.text for para in doc.paragraphs if para.text.strip()])
        return [text]  # Return as a list for consistency

    def loade(self, file_path):
        documents = ''
        if file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
        elif file_path.endswith('.docx'):
            loader = Docx2txtLoader(file_path)
        else:
            raise ValueError(
                "Unsupported file format. Please provide a PDF or Word document.")
        return loader

    def embedd(self, file_path):
        # loader = PyPDFLoader(file_path)
        # documents = loader.load()

        loader = self.loade(file_path)
        documents = loader.load()

        if (file_path.endswith('.docx')):
            documents = [Document(page_content=doc) if isinstance(
                doc, str) else doc for doc in documents]

        chunks = self.splitter.split_documents(documents)
        chunks = [doc.page_content for doc in chunks]
        vectors = self.embedingAPI.embed_documents(chunks)
        return (chunks, vectors)

    def embedquerry(self, querry):
        querry = self.embedingAPI.embed_query(querry)
        return querry

    def SaveCsv(self, file_path, name, vectors, chunks):
        df = pd.DataFrame({
            "chunks": chunks,
            "vectors": vectors
        })
        if not file_path.endswith("/"):
            file_path += "/"

        name = name[:name.index('.')]
        locat = file_path+name+'.csv'
        df.to_csv(locat, index=False)

    def ReadFromFile(self, file_path):
        df = pd.read_csv(file_path)
        chunks = df["chunks"].tolist()
        vectors = df["vectors"].apply(
            ast.literal_eval).tolist()  # Convert strings to lists
        return (chunks, vectors)

    def cosine_search(self, vectors, query_vector):
        vectors = np.array(vectors)
        query_vector = np.array(query_vector)
        query_vector = query_vector.reshape(1, -1)
        distances = cosine_similarity(vectors, query_vector)
        closest_index = np.argmax(distances)
        return closest_index

    def similarity_percentage(self, x, y):
        xn = min([x, y])
        yn = max([x, y])
        return (1 - abs(xn - yn) / abs(yn)) * 100

    def cosine_search_top3(self, vectors, query_vector, threshold=20):
        top_3_indices = []
        vectors = np.array(vectors)
        query_vector = np.array(query_vector)
        query_vector = query_vector.reshape(1, -1)

        distances = cosine_similarity(vectors, query_vector)

        similarities = distances.flatten() * 100

        closest_index = np.argmax(distances)
        if (similarities[closest_index] < threshold):
            return ([], 0)
        sorted_indices = np.argsort(similarities)[::-1]

        for i in range(0, len(sorted_indices[:3])):
            if (similarities[sorted_indices[i]] >= threshold):
                top_3_indices.append(sorted_indices[i])

        similarities_score = 0
        for i in top_3_indices:
            similarities_score += distances[i]
        return (top_3_indices, similarities_score)

    def cosine_search_top3_t(self, vectors, query_vector, threshold=20):
        vectors = np.array(vectors)
        query_vector = np.array(query_vector)
        query_vector = query_vector.reshape(1, -1)

        distances = cosine_similarity(vectors, query_vector)
        closest_index = np.argmax(distances)

        similarities = distances.flatten() * 100
        print(similarities)
        # Reverse for descending order
        sorted_indices = np.argsort(similarities)[::-1]
        top_3_indices = [
            i for i in sorted_indices if self.similarity_percentage(similarities[i], similarities[closest_index]) >= threshold][:3]
        # print(top_3_vectors, " top 3 vectors")
        # print("-"*20)
        # print(top_3_similarities, " top 3 similierts")
        # print("-"*20)
        similarities_score = 0
        for i in top_3_indices:
            similarities_score += distances[i]
        return (top_3_indices, similarities_score)

    def cosine_search_chunks(self, data, query_vector):
        chunks = data[0]
        vectors = np.array(data[1])
        query_vector = np.array(query_vector)
        query_vector = query_vector.reshape(1, -1)
        distances = cosine_similarity(vectors, query_vector)
        closest_index = np.argmax(distances)
        return chunks[closest_index]


# ------------------------------------------------------------------------------------------------

    def Vectoraiz(self, file_path):
        self.file_contet = PyPDFLoader(file_path).load()
        docs = self.splitter.split_documents(self.file_contet)
        self.db = Chroma.from_documents(docs, OpenAIEmbeddings())
        return self.db

    def querry(self, question: str):
        vectorEmbedQuery = OpenAIEmbeddings().embed_query(question)
        answer = self.db.similarity_search_by_vector(vectorEmbedQuery)
        return answer
# ------------------------------------------------------------------------------------------------


if __name__ == "__main__":
    num_samples = 100
    dimensionality = 5

    data = np.random.rand(num_samples, dimensionality)
    query_vector = np.random.rand(dimensionality)

# Calculate cosine similarities between the query vector and the dataset
    similarities = cosine_similarity(data, [query_vector])

# Find the most similar vector
    most_similar_index = np.argmax(similarities)
    most_similar_vector = data[most_similar_index]

    my_list = [42, 73]
    last_10_items = my_list[-10:]
    print(last_10_items)

# Calculate Euclidean distances between the query vector and the dataset

# Find the closest vector

# for i in answer:

# for i in answer:
# for i in answer:
# print(i.page_content)
# print("-----------------")

# PandasSave(vectorEmbedQuery)
# a = ReadFromFile('vector2.csv')
# for i in answer2:
#    print(i.page_content)
#    print('1')
# print("-----------------")

# print(a)
# If it's a list or numpy array
