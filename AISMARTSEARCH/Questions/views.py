import ast
import collections
import json
import os
import re
import shutil
from itertools import chain
from os import walk
from pathlib import Path
from urllib.parse import unquote

import numpy as np
import tiktoken
from django.conf import settings
from django.contrib.auth.models import User
from django.core.files.storage import FileSystemStorage
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from FileServis import FileServices
from Main import Parsers
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from .forms import FileUploadForm, MakeDirForm, UserValueForm
from .models import Chunk, ContractState, History, UserValues


def clear_all_contract_fields():
    contract_states = ContractState.objects.all()
    count = contract_states.count()

    if count == 0:
        print("No ContractState records found.")
        return

    for cs in contract_states:
        cs.fields = ""  # clear the fields
        cs.save()

    print(f"Cleared 'fields' for {count} ContractState records.")

    from django.http import JsonResponse


def check_messages_length_approx(messages):
    total_chars = sum(len(msg.get("content", "")) for msg in messages)
    total_tokens_estimate = total_chars // 4  # rough estimate
    return {
        "num_messages": len(messages),
        "total_chars": total_chars,
        "estimated_tokens": total_tokens_estimate,
    }


def print_all_users(request):
    username = "qerimqerimAi"
    password = "hello123"

    # Delete user (and ContractState via CASCADE)
    try:
        user = User.objects.get(username=username)
        user.delete()
        print(f"Deleted user {username} and related objects.")
    except User.DoesNotExist:
        print(f"User {username} did not exist, skipping delete.")

    # Recreate user
    new_user = User.objects.create_user(username=username, password=password)
    ContractState.objects.create(user=new_user)
    print(f"Recreated user {username} with password {password}")

    # Get all users
    users = User.objects.all()
    if not users:
        print("No users found.")
        return JsonResponse({"1": []})

    print("All Django users:")
    usernames = []
    for user in users:
        print(user)
        usernames.append(user.username)
        print(
            f"ID: {user.id}, Username: {user.username}, Email: {user.email}, Is Staff: {user.is_staff}"
        )

    return JsonResponse({"1": usernames})


def getalldirs(mypath):
    user = []
    for dirpath, dirnames, filenames in walk(mypath):
        user = dirnames
        break
    return user


@csrf_exempt
def fileuploadfront(request, user, fileType):

    if request.method == "POST" and request.FILES.get("file"):

        path = settings.STATIC_UPLOAD_DIR
        if fileType == "akt":
            path = settings.INTERNAL_ACT

        uploaded_file = request.FILES["file"]

        user_dir = Path(path) / user
        user_dir.mkdir(
            parents=True, exist_ok=True
        )  # Create the user's directory if it doesn't exist

        file_path = user_dir / uploaded_file.name

        if file_path.exists():
            return JsonResponse({"success": False, "message": "The file already exists."})
        my_file = Path(f"{path}/{user}/{uploaded_file.name}")
        print(f"this is the file path {my_file}")

        if not my_file.is_file():
            if fileType == "normal":
                save_file(uploaded_file, user)
            else:
                save_file_akt(uploaded_file, user)

        with open(file_path, "wb+") as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        return JsonResponse(
            {
                "success": True,
                "message": "hey you have uploaded it correctly",
                "filename": uploaded_file.name,
            }
        )

    return JsonResponse(
        {"success": False, "message": "No file uploaded or wrong method"}, status=400
    )


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_user_values(request):
    try:
        uv = UserValues.objects.get(user=request.user)
        return Response(UserValuesSerializer(uv).data)
    except UserValues.DoesNotExist:
        return Response({"error": "User values not found"}, status=404)


@api_view(["POST"])
def signup(request):
    username = request.data.get("username")
    password = request.data.get("password")
    if User.objects.filter(username=username).exists():
        return Response({"error": "Username already exists"}, status=400)
    user = User.objects.create_user(username=username, password=password)

    ContractState.objects.create(user=user)

    UserValues.objects.create(
        user=username,
        splitter="CharacterTextSplitter",  # default
        chunksize=500,
        overlap=50,
        temp=0.5,
    )

    """
    TODO: make it add internal akt
    """
    makedir(username)
    makedir_internal_akt(username)
    makedir_genert(username)
    return Response({"message": "User created successfully"}, status=201)


def allFileformat(mypath, format):
    files = []
    for dirpath, dirnames, filenames in walk(mypath):
        csv_files = [file for file in filenames if file.endswith(format)]
        files.extend(csv_files)
        break
    return files


def remove_csv(user):
    user_path = settings.STATIC_UPLOAD_DIR + f"/{user}"

    files = allFileformat(user_path, ".csv")

    for i in files:
        newpath = user_path + f"/{i}"
        if os.path.exists(newpath):
            os.remove(newpath)
            print(f"File '{newpath}' has been deleted.")
        else:
            print(f"File '{newpath}' does not exist.")


def recrate_csvs(user_path, user, parser):
    pdf_files = allFileformat(user_path, ".pdf")
    word_files = allFileformat(user_path, ".docx")
    files = pdf_files + word_files
    for file_name in files:
        file_url = f"{settings.STATIC_UPLOAD_DIR}/{user}/{file_name}"
        fileChunks, fileEmbedings = parser.embedd(file_url)
        parser.SaveCsv(
            settings.STATIC_UPLOAD_DIR + "/" + user,
            filename.rsplit(".", 1)[0],
            fileEmbedings,
            fileChunks,
        )


def recrate_contract_template(request):

    parser = Parsers(settings.OPENAI_KEY)
    cr_path = os.path.join(settings.BASE_DIR, "hi")
    pdf_files = allFileformat(cr_path, ".pdf")
    word_files = allFileformat(cr_path, ".docx")

    print("this is da path ", cr_path)
    files = pdf_files + word_files
    for file_name in files:
        file_url = f"{cr_path}/{file_name}"
        fileChunks, fileEmbedings = parser.embedd(file_url)
        parser.SaveCsv(
            settings.CONTRACT_TEMPALTES,
            file_name[: file_name.rindex(".")],
            fileEmbedings,
            fileChunks,
        )
    return JsonResponse({"mes": "val"})


def reembedfiles(user):
    user_path = settings.STATIC_UPLOAD_DIR + f"/{user}"
    user_value = UserValues.objects.filter(user=user).first()
    parser = Parsers(settings.OPENAI_KEY)
    spliter = user_value.splitter
    chunksize = user_value.chunksize
    overlap = user_value.overlap
    parser.SetSpliter(spliter=spliter, chuncksize=chunksize, overlap=overlap)

    remove_csv(user)
    recrate_csvs(user_path, user, parser)


def addfiledata(dic, file_name, chunks, vectors, sim_score):
    dic[file_name] = {"chunks": chunks, "vectors": vectors, "score": sim_score}


def sort_data(files_data):
    sorted_files = sorted(files_data.items(),
                          key=lambda item: item[1]["score"], reverse=True)
    top_10_files = sorted_files[:10]
    top_10_chunks = list(chain.from_iterable(
        item[1]["chunks"] for item in top_10_files))

    top_10_vectors = list(chain.from_iterable(
        item[1]["vectors"] for item in top_10_files))

    sorted_files_dict = collections.OrderedDict(top_10_files)

    return (top_10_chunks, top_10_vectors, sorted_files_dict)


def system_file_parserT(querry_vector, user, base_path=settings.STATIC_UPLOAD_DIR):
    mypath = base_path + "/" + user + "/"
    parser = Parsers(settings.OPENAI_KEY)  # ✅

    vectorlist = []
    chunkslist = []
    files = allFileformat(mypath, ".csv")
    files_data = {}
    for i in files:
        chunks, vectors = parser.ReadFromFile(mypath + i)
        # closest_index = parser.cosine_search(vectors, querry_vector)
        top3, similariti_score = parser.cosine_search_top3(
            vectors, querry_vector, 40)
        for j in top3:
            chunkslist.append(chunks[j])
            vectorlist.append(vectors[j])
        if len(chunkslist) > 0:
            addfiledata(files_data, i, chunkslist,
                        vectorlist, similariti_score)

            # print(files_data[i]['chunks'], files_data[i]['score'])

        chunkslist = []
        vectorlist = []

    top_10_chunks, top_10_vectors, sorted_files_dict = sort_data(files_data)
    return (top_10_chunks, top_10_vectors, sorted_files_dict)


def system_file_parser_law(querry_vector, user, base_path=settings.BASE_LAWS):
    mypath = base_path + "/"
    parser = Parsers(settings.OPENAI_KEY)  # ✅

    vectorlist = []
    chunkslist = []
    files = allFileformat(mypath, ".csv")
    files_data = {}
    for i in files:
        chunks, vectors = parser.ReadFromFile(mypath + i)
        # closest_index = parser.cosine_search(vectors, querry_vector)
        top3, similariti_score = parser.cosine_search_top3(
            vectors, querry_vector, 30)
        for j in top3:
            chunkslist.append(chunks[j])
            vectorlist.append(vectors[j])
        if len(chunkslist) > 0:
            addfiledata(files_data, i, chunkslist,
                        vectorlist, similariti_score)

            # print(files_data[i]['chunks'], files_data[i]['score'])

        chunkslist = []
        vectorlist = []

    top_10_chunks, top_10_vectors, sorted_files_dict = sort_data(files_data)
    return (top_10_chunks, top_10_vectors, sorted_files_dict)


def system_file_parser(query_vectors, user, base_path=settings.STATIC_UPLOAD_DIR):

    mypath = base_path + "/" + user + "/"
    parser = Parsers(settings.OPENAI_KEY)

    files = allFileformat(mypath, ".csv")
    files_data = {}

    for file in files:
        chunks, vectors = parser.ReadFromFile(mypath + file)
        chunkslist = []
        vectorlist = []
        total_score = 0

        # Loop over each vector from legal context
        for law_vector in query_vectors:
            top3, similarity_score = parser.cosine_search_top3(
                vectors, law_vector, 30)
            for idx in top3:
                chunkslist.append(chunks[idx])
                vectorlist.append(vectors[idx])
            total_score += similarity_score  # accumulate relevance

        if chunkslist:
            addfiledata(files_data, file, chunkslist, vectorlist, total_score)

    top_10_chunks, top_10_vectors, sorted_files_dict = sort_data(files_data)
    return top_10_chunks, top_10_vectors, sorted_files_dict


def get_case_tamplates(query_vectors):
    mypath = settings.CONTRACT_TEMPALTES + "/"
    parser = Parsers(settings.OPENAI_KEY)

    files = allFileformat(mypath, ".csv")
    print("this is case tempalte" * 19)
    print(files)
    print("this is case tempalte" * 19)

    files_data = {}

    # Determine expected vector length from query_vector
    expected_len = len(query_vectors)

    for file in files:
        chunks, vectors = parser.ReadFromFile(mypath + file)

        # Filter out vectors that do not match the expected length
        filtered_chunks = []
        filtered_vectors = []
        for c, v in zip(chunks, vectors):
            if isinstance(v, list) and len(v) == expected_len:
                filtered_chunks.append(c)
                filtered_vectors.append(v)
            else:
                print(
                    f"Skipping vector with invalid length in file {file}: {v}")

        if not filtered_vectors:
            continue  # skip this file if no valid vectors

        chunkslist = []
        vectorlist = []
        total_score = 0

        # Compute top3 similarity with filtered vectors
        top3, similarity_score = parser.cosine_search_top3(
            filtered_vectors, query_vectors, 30)

        for idx in top3:
            chunkslist.append(filtered_chunks[idx])
            vectorlist.append(filtered_vectors[idx])
            total_score += similarity_score  # accumulate relevance

        if chunkslist:
            addfiledata(files_data, file, chunkslist, vectorlist, total_score)

    top_10_chunks, top_10_vectors, sorted_files_dict = sort_data(files_data)
    return top_10_chunks, top_10_vectors, sorted_files_dict


def get_case_tamplates2(query_vectors):
    mypath = settings.CONTRACT_TEMPALTES + "/"
    parser = Parsers(settings.OPENAI_KEY)

    files = allFileformat(mypath, ".csv")
    print("this is case tempalte" * 19)
    print(files)
    print("this is case tempalte" * 19)
    files_data = {}
    for file in files:
        chunks, vectors = parser.ReadFromFile(mypath + file)
        chunkslist = []
        vectorlist = []
        total_score = 0

        # Loop over each vector from legal context
        top3, similarity_score = parser.cosine_search_top3(
            vectors, query_vectors, 30)
        for idx in top3:
            chunkslist.append(chunks[idx])
            vectorlist.append(vectors[idx])
            total_score += similarity_score  # accumulate relevance

        if chunkslist:
            addfiledata(files_data, file, chunkslist, vectorlist, total_score)

    top_10_chunks, top_10_vectors, sorted_files_dict = sort_data(files_data)
    return top_10_chunks, top_10_vectors, sorted_files_dict


def get_case_files(query_vectors):
    parser = Parsers(settings.OPENAI_KEY)  # ✅
    mypath = settings.BASE_LAWS + "/"
    files = get_all_base_laws()

    files_data = {}
    for file in files:
        chunks, vectors = parser.ReadFromFile(mypath + file)
        chunkslist = []
        vectorlist = []
        total_score = 0

        # Loop over each vector from legal context
        for law_vector in query_vectors:
            top3, similarity_score = parser.cosine_search_top3(
                vectors, law_vector, 30)
            for idx in top3:
                chunkslist.append(chunks[idx])
                vectorlist.append(vectors[idx])
            total_score += similarity_score  # accumulate relevance

        if chunkslist:
            addfiledata(files_data, file, chunkslist, vectorlist, total_score)

    top_10_chunks, top_10_vectors, sorted_files_dict = sort_data(files_data)
    return top_10_chunks, top_10_vectors, sorted_files_dict


"""
YO METIII  sooo tomore with it beaing the 29 we will bve adding multiepr layers to this what do i mean is that the refrec types of it
will be mutlitle insted of 1 specifike file like it curralty is to multiple ones sooo yes
lets do it darlings
"""


def Get_aktin_i_brendshem_data(querry_vector):
    parser = Parsers(settings.OPENAI_KEY)  # ✅
    top_law_chunks, top_law_vector = parser.ReadFromFile(
        settings.BASE_LAWS + "/" +
        "Draft Akti i Brendshëm - Kosove (7) (1).csv"
    )

    top3, similariti_score = parser.cosine_search_top3(
        top_law_vector, querry_vector, 30)
    chunkslist = []
    vectorlist = []
    for j in top3:
        chunkslist.append(top_law_chunks[j])
        vectorlist.append(top_law_vector[j])

    return (chunkslist, vectorlist)


def Get_law_data(querry_vector):
    parser = Parsers(settings.OPENAI_KEY)  # ✅
    top_law_chunks, top_law_vector = parser.ReadFromFile(
        settings.BASE_LAWS + "/" + "LIGJI__NR.csv")

    top3, similariti_score = parser.cosine_search_top3(
        top_law_vector, querry_vector, 30)
    chunkslist = []
    vectorlist = []
    for j in top3:
        chunkslist.append(top_law_chunks[j])
        vectorlist.append(top_law_vector[j])

    return (chunkslist, vectorlist)


def addReview_one(user, filename, messages):
    """
    data: list of tuples (file_name: str, content: str)
    messages: list of dicts representing the existing messages list to which system entries will be appended

    For each file and each chunk inside its content, add a system message describing the chunk.
    """
    parser = Parsers(settings.OPENAI_KEY)  # ✅

    userpath = os.path.join(settings.STATIC_UPLOAD_DIR, user)
    file = filename + ".csv"
    data = {}
    chunks, vectores = parser.ReadFromFile(os.path.join(userpath, file))
    # Combine all chunks into a single string
    long_text = "\n\n".join(chunks)
    data[file] = long_text
    for file_name, chunk in data.items():
        new_entry = {
            "role": "system",
            "content": (
                f"This is part of the document '{file_name}' relevant for legal review (Section {file}):\n{chunk}"
            ),
        }
        messages.append(new_entry)

    print(f"relted fiels ❗{file}")


""" """


def chunk_list(lst, n):
    """Split list into chunks of size n (last batch may be smaller)."""
    for i in range(0, len(lst), n):
        yield lst[i: i + n]


def review_in_batches(user, query, client):
    print("LAW review started ⚠️")
    userpath = os.path.join(settings.STATIC_UPLOAD_DIR, user)
    files = allFileformat(userpath, ".csv")

    print(files)

    responses = []

    # Split files into groups of 3
    for idx, file_group in enumerate(chunk_list(files, 3), start=1):
        print(f"➡️ Starting review for batch {idx}: {file_group}")

        # fresh messages per group
        messages2 = [
            {
                "role": "system",
                "content": (
                    "You are a highly skilled legal assistant specializing in legal document review.\n\n"
                    "Your task is to review all provided legal documents carefully and return a clear, structured legal analysis. Focus specifically on:\n\n"
                    "- Identifying potential legal risks or liabilities\n"
                    "- Noting missing clauses or provisions\n"
                    "- Highlighting compliance issues or inconsistencies\n"
                    "- Pointing out ambiguous or vague language\n"
                    "- Detecting contradictions across documents\n\n"
                    "If the user hasn’t asked for a specific focus, perform a general legal review. "
                    "If a review scope is provided (e.g., employment law, data privacy), prioritize findings related to that.\n\n"
                    "Your tone should be professional, neutral, and clear. Use bullet points, sections, or numbered lists where appropriate.\n\n"
                    "⚠️ Do **not** summarize the documents unless explicitly asked.\n"
                    "⚠️ Do **not** provide general legal theory — focus on what's present or missing in the provided text.\n\n"
                    "If something is unclear, flag it. If everything looks good, say so explicitly.\n\n"
                    "Example output structure:\n"
                    "-------------------------\n"
                    "Legal Review:\n\n"
                    "1. Missing Clauses\n"
                    "   - No termination clause detected.\n"
                    "   - Lacks confidentiality agreement section.\n\n"
                    "2. Inconsistencies\n"
                    "   - Section 4 conflicts with Section 2 regarding payment terms.\n\n"
                    "3. Risk Assessment\n"
                    "   - Clause 8 creates a potential liability due to vague language.\n\n"
                    "   - Specify if you can the Level of the risk Low/Medium/High\n"
                    "4. General Observations\n"
                    "   - Formatting is inconsistent.\n"
                    "   - Passive voice used in critical obligations.\n\n"
                    "End your review with a one-line summary of overall document quality "
                    "(e.g., “The document is mostly complete but contains moderate risk areas.”)"
                ),
            },
            {
                "role": "system",
                "content": (
                    "The following guidance document outlines the standards and best practices for legal drafting and review in Kosovo. "
                    "Refer to it strictly in all assessments:\n\n"
                    f"{add_guidelines()}"
                ),
            },
            {"role": "user", "content": f"Current Question: {query}"},
        ]

        # add just these 3 files
        addReviewContext_batch(user, file_group, messages2)

        # one GPT call per group
        res = client.chat.completions.create(
            model="gpt-4o",
            messages=messages2,
            max_tokens=1500,
            temperature=0.2,
        )

        response_message = res.choices[0].message
        responses.append({"batch": idx, "files": file_group,
                         "review": response_message.content})

    return responses


def addReviewContext_batch(user, file_group, messages):
    """Add review context for a specific group of files."""
    parser = Parsers(settings.OPENAI_KEY)
    userpath = os.path.join(settings.STATIC_UPLOAD_DIR, user)

    for file in file_group:
        chunks, _ = parser.ReadFromFile(os.path.join(userpath, file))
        long_text = "\n\n".join(chunks)

        new_entry = {
            "role": "system",
            "content": (
                f"This is part of the document '{file}' relevant for legal review:\n{long_text}"
            ),
        }
        messages.append(new_entry)

        print(f"related files ❗ {file_group}")


""" """


def addReviewContext(user, messages):
    """
    data: list of tuples (file_name: str, content: str)
    messages: list of dicts representing the existing messages list to which system entries will be appended

    For each file and each chunk inside its content, add a system message describing the chunk.
    """
    parser = Parsers(settings.OPENAI_KEY)  # ✅

    userpath = os.path.join(settings.STATIC_UPLOAD_DIR, user)
    files = allFileformat(userpath, ".csv")
    data = {}
    for i in files:

        chunks, vectores = parser.ReadFromFile(os.path.join(userpath, i))
        # Combine all chunks into a single string
        long_text = "\n\n".join(chunks)
        data[i] = long_text

    for file_name, chunk in data.items():
        new_entry = {
            "role": "system",
            "content": (
                f"This is part of the document '{file_name}' relevant for legal review (Section {i}):\n{chunk}"
            ),
        }
        messages.append(new_entry)

    print(f"relted fiels ❗{files}")


def addCaseFileContext(case_chunks, messages):
    chunks = "\n".join([f"```text\n{chunk}\n```" for chunk in case_chunks])
    case_entry = {
        "role": "assistant",
        "content": (
            "The following excerpts are from case files that may be relevant to the current legal question. "
            "Review them carefully and apply them where applicable in your response:\n"
            f"{chunks}"
        ),
    }
    messages.append(case_entry)


def addLawContext(law_chunks, message):
    chunks = "\n".join([f"```text\n{chunk}\n```" for chunk in law_chunks])
    contract_entry = {
        "role": "system",
        "content": (
            "The following contract or legal clauses contain key information relevant to the current legal question. "
            "Analyze them for risks, inconsistencies, or issues:\n"
            f"{chunks}"
        ),
    }
    message.append(contract_entry)


def addInternalAkteContext(akte_chunks, message):
    chunks = "\n".join([f"```text\n{chunk}\n```" for chunk in akte_chunks])
    akte_entry = {
        "role": "system",
        "content": (
            "The following internal aktes or company policies contain key information relevant to the current legal or procedural question. "
            "Analyze them for compliance, clarity, risks, or inconsistencies:\n"
            f"{chunks}"
        ),
    }
    message.append(akte_entry)


def addContreact(data, message):
    for file_name, content in data:
        # print(file_name + '-' * 20)

        chunks = content

        newdic = {
            "role": "system",
            "content": (
                f"The following file: '{file_name}' to generate a contrect: \n" f"{chunks}"
            ),
        }
        message.append(newdic)


def addContext(data, message):
    for file_name, content in data.items():
        # print(file_name + '-' * 20)

        chunks = "\n".join(
            [f"```text\n{chunk}\n```" for chunk in content["chunks"]])

        newdic = {
            "role": "system",
            "content": (
                f"The following file: '{
                    file_name}' contains relevant information to support the answer: \n"
                f"{chunks}"
            ),
        }
        message.append(newdic)


def addHistory(question_history, answer_history, message):
    question_history = question_history[-10:]
    answer_history = answer_history[-10:]

    # Combine questions and answers into the message
    for index, (q, a) in enumerate(zip(question_history, answer_history)):
        question_entry = {"role": "user", "content": f"Past Question : {q}"}
        answer_entry = {"role": "assistant", "content": f"Past Answer : {a}"}
        message.append(question_entry)
        message.append(answer_entry)


def gettemp(user):
    return UserValues.objects.filter(user=user).first().temp


def context_aware_responsesT(query, Question_history, Answer_history, data, user):
    # openai.api_key = settings.OPENAI_K
    temp = gettemp(user)
    client = OpenAI(api_key=settings.OPENAI_KEY)
    # print(query, "this is querry", "-"*100)
    # if (len(data) == 0):
    #    return "The context does not contain sufficient information to answer this question.--2"
    messages = [
        {
            "role": "system",
            "content": (
                "You are an AI Legal Advisor specialized in providing legal assistance and advice based strictly on the provided context. "
                "Your behavior and responses must follow these rules:\n\n"
                "1. **Use Only Provided Sources**:\n"
                "- You must answer questions solely based on the given context, provided documents, or user-supplied information.\n"
                "- You are not permitted to use external knowledge, personal opinions, or make assumptions beyond the provided material.\n\n"
                "2. **Legal-Focused and Contextual Responses**:\n"
                "- Deliver advice, clarifications, or interpretations as a legal professional would, while strictly adhering to the available information.\n\n"
                "3. **Clear, Precise, and Formal Language**:\n"
                "- Provide legally precise, clear, and concise responses.\n"
                "- Avoid unnecessary elaborations unless specifically requested by the user.\n\n"
                "- Make it as short and consisnet as you can baste on the relevt inpout. \n\n"
                "4. **Professional and Respectful Tone**:\n"
                "- Maintain a formal, respectful, and professional tone appropriate for legal communications.\n\n"
                "5. **Handle Uncertainties Transparently**:\n"
                "- If the context is unclear, incomplete, or contradictory, explicitly acknowledge the limitation and request clarification or note that a conclusion cannot be drawn.\n\n"
                "You will now begin assisting the user with legal questions based solely on the provided sources and context."
            ),
        },
        {"role": "user", "content": f"Current Question: {query}"},
    ]

    # addHistory(Question_history, Answer_history, messages)
    addContext(data, messages)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        max_tokens=900,  # test on the higer end to the lowest 50-600
        temperature=temp,  # Strict and deterministic responses
    )
    # print(messages)
    response_message = response.choices[0].message.content
    return response_message


def genert_tools():
    tools = [
        {
            "type": "function",
            "function": {
                "name": "trigger_contract_generation",
                "description": "Call this if the user's message is about generating a legal contract or document.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "reason": {
                            "type": "string",
                            "description": "Short explanation of why contract generation is being triggered based on the user’s question.",
                        }
                    },
                    "required": ["reason"],
                },
            },
        },
        {
            "type": "function",
            "function": {
                "name": "trigger_legal_review",
                "description": "Call this if the user's message indicates they want a legal review of their documents.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "reason": {
                            "type": "string",
                            "description": "Short explanation of why legal review is being triggered based on the user’s question.",
                        },
                        "files_to_review": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "List of filenames or document IDs to include in the legal review. If empty or omitted, all files are reviewed.",
                        },
                    },
                    "required": ["reason"],
                },
            },
        },
    ]
    return tools


def add_guidelines():
    parser = Parsers(settings.OPENAI_KEY)
    all_chunks = []

    review_dir = settings.LEAGLEREVIEW

    long_text = ""
    for i in allFileformat(review_dir, ".csv"):
        chunks, vec = parser.ReadFromFile(settings.LEAGLEREVIEW + "/" + i)

        long_text = "\n\n".join(chunks)
    print("hello this is it long text hope it works woohoohoho ❤️")
    print(long_text[:150])
    return long_text


def convert_ndarray_to_list(obj):
    """
    Recursively convert all numpy arrays to lists in a nested structure.
    """
    if isinstance(obj, dict):
        return {k: convert_ndarray_to_list(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_ndarray_to_list(i) for i in obj]
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    else:
        return obj


def context_aware_responses(query, law_data, ak_data, case_data, data, user):
    # openai.api_key = settings.OPENAI_K
    temp = gettemp(user)
    client = OpenAI(api_key=settings.OPENAI_KEY)
    # print(query, "this is querry", "-"*100)
    # if (len(data) == 0):
    #    return "The context does not contain sufficient information to answer this question.--2"
    messages = [
        {
            "role": "system",
            "content": (
                "You are an AI Legal Advisor. You must answer **only based on the provided case file and legal context**. "
                "Follow these priorities:\n\n"
                "1. If the user asks a direct legal question about their situation → respond based only on the relevant case file context. "
                "   - Summarize the facts provided.\n"
                "   - Apply the law context strictly to those facts.\n"
                "   - Keep the answer short, precise, and professional.\n\n"
                "2. If the question is not clearly answered by the context → explicitly state what information is missing and ask the user for clarification.\n\n"
                "3. Do **not** provide general explanations of company policies or unrelated laws unless they are directly relevant.\n\n"
                "Tone: professional, clear, formal. Avoid filler, avoid repeating irrelevant policies."
            ),
        },
        {"role": "user", "content": f"Current Question: {query}"},
    ]

    """
    THIS SHULD insted of add lawcontext add case file context it should be a gpt deffing respose baste on the relvet question
    and if the users ask fora  review okay you ask for a review if i ask for lawconte
    i am giving it  bastes on just 1 thign static infromaot maybe the question is not reletd to interlaakte
    but still give  witch would give the network more ireelevet information to fillter by
    soo this sohuld be a tool top messesg is a base for the system
    """
    # addHistory(Question_history, Answer_history, messages)

    addLawContext(law_data, messages)
    # addCaseFileContext(case_data, messages)
    addInternalAkteContext(ak_data, messages)

    addContext(data, messages)

    print("length ->", check_messages_length_approx(messages))

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        max_tokens=300,  # test on the higer end to the lowest 50-600
        temperature=temp,  # Strict and deterministic responses
        tools=genert_tools(),
    )
    # print(messages)
    response_message = response.choices[0].message
    if response_message.tool_calls:
        for tool_call in response_message.tool_calls:
            if tool_call.function.name == "trigger_contract_generation":
                print("⚠️ Contract generation triggered!")
                reason = json.loads(tool_call.function.arguments)["reason"]
                parser = Parsers(settings.OPENAI_KEY)
                query_vector = parser.embedquerry(query)
                "just fix this :TODO i need ot fix this fr thos so ecetyinhg worsk just fine "
                chunklis, vectorlis, files = get_case_tamplates(query_vector)
                messages2 = [
                    {
                        "role": "system",
                        "content": (
                            "You are a legal assistant that helps generate contracts. "
                            "From the given context, identify what fields need to be filled by the user to complete the contract (e.g., names, dates, fees). "
                            "Remove any duplicates fileds "
                            "Return them as a comma-separated list with no explanation. Only list the missing field names."
                        ),
                    },
                    {"role": "user", "content": f"Current Question: {query}"},
                ]
                print("this are relevt files ")
                print(files)
                print("this is for loop")
                allstring = []

                for i in files:
                    contract_state, _ = ContractState.objects.get_or_create(
                        user=getuser(user))

                    # serialize the dict to a JSON string
                    safe_data = convert_ndarray_to_list(
                        files[i])  # convert all ndarrays to lists

                    contract_state.file = json.dumps(
                        safe_data, ensure_ascii=False)
                    contract_state.save()

                    print("<---3, ", i)
                    chu, vec = parser.ReadFromFile(
                        settings.CONTRACT_TEMPALTES + "/" + i)
                    allstring.append((i, " ".join(chu)))
                    break

                print("this are relevt files ")
                addContreact(allstring, messages2)
                res = client.chat.completions.create(
                    model="gpt-4o",
                    messages=messages2,
                    max_tokens=300,  # test on the higer end to the lowest 50-600
                    temperature=temp,  # Strict and deterministic responses
                )

                response_message = res.choices[0].message
                required_fields = res.choices[0].message.content.strip().split(
                    ",")
                required_fields = [field.strip()
                                   for field in required_fields if field.strip()]

                # if len(required_fields) > 0:
                #    settings.USER_INFO = True
                #    settings.FIELDS = "\n".join(f"- {field}" for field in required_fields)
                if len(required_fields) > 0:
                    contract_state, _ = ContractState.objects.get_or_create(
                        user=getuser(user))
                    contract_state.user_info = True
                    contract_state.fields = "\n".join(
                        f"- {field}" for field in required_fields)
                    contract_state.save()

                print(required_fields)

                print("L)@@@" * 2)
                print(response_message, " ", response_message.content)
                return (response_message.content, "Contract_generted")
                # return {"action": "generate_contract", "reason": reason, "triggered": True}

            if tool_call.function.name == "trigger_legal_review":

                response_message = review_in_batches(user, query, client)
                return (response_message, "Review")
                # if tool_call.function.name == "trigger_legal_review":
                # messages2 = [
                # {
                # "role": "system",
                # "content": (
                # "You are a highly skilled legal assistant specializing in legal document review.\n\n"
                # "Your task is to review all provided legal documents carefully and return a clear, structured legal analysis. Focus specifically on:\n\n"
                # "- Identifying potential legal risks or liabilities\n"
                # "- Noting missing clauses or provisions\n"
                # "- Highlighting compliance issues or inconsistencies\n"
                # "- Pointing out ambiguous or vague language\n"
                # "- Detecting contradictions across documents\n\n"
                # "If the user hasn’t asked for a specific focus, perform a general legal review. "
                # "If a review scope is provided (e.g., employment law, data privacy), prioritize findings related to that.\n\n"
                # "Your tone should be professional, neutral, and clear. Use bullet points, sections, or numbered lists where appropriate.\n\n"
                # "⚠️ Do **not** summarize the documents unless explicitly asked.\n"
                # "⚠️ Do **not** provide general legal theory — focus on what's present or missing in the provided text.\n\n"
                # "If something is unclear, flag it. If everything looks good, say so explicitly.\n\n"
                # "You do not need to ask follow-up questions unless the prompt requests clarification.\n\n"
                # "Example output structure:\n"
                # "-------------------------\n"
                # "Legal Review:\n\n"
                # "1. Missing Clauses\n"
                # "   - No termination clause detected.\n"
                # "   - Lacks confidentiality agreement section.\n\n"
                # "   - if somthing can be change specifke it and the valid change base on it \n"
                # "2. Inconsistencies\n"
                # "   - Section 4 conflicts with Section 2 regarding payment terms.\n\n"
                # "3. Risk Assessment\n"
                # "   - Clause 8 creates a potential liability due to vague language.\n\n"
                # "   - specife if you can the Level of the risk Low/Medium/High"
                # "4. General Observations\n"
                # "   - Formatting is inconsistent.\n"
                # "   - Passive voice used in critical obligations.\n\n"
                # "End your review with a one-line summary of overall document quality "
            # "(e.g., “The document is mostly complete but contains moderate risk areas.”)"
            # ),
            # },
            # {"role": "user", "content": f"Current Question: {query}"},
            # ]
            # messages2.insert(
            # 1,
            # {
            # "role": "system",
            # "content": (
            # "The following guidance document outlines the standards and best practices for legal drafting and review in Kosovo. "
            # "Refer to it strictly in all assessments:\n\n"
            # <- load this from the PDF or pre-parsed text
            # f"{add_guidelines()}"
            # ),
            # },
            # )

            # print("LAW reveiw stared ⚠️")
            # userpath = os.path.join(settings.STATIC_UPLOAD_DIR, user)
            # files = allFileformat(userpath, ".csv")

            # print(files)

            # addReviewContext(user, messages2)
            # res = client.chat.completions.create(
            # model="gpt-4o",
            # messages=messages2,
            # max_tokens=1500,  # test on the higer end to the lowest 50-600
            # temperature=0.2,  # Strict and deterministic responses
            # )

            # response_message = res.choices[0].message
            # return (response_message.content, "Review")
    # If no function triggered, continue normally
    return (response_message.content, "Normal")


def unpack_history(history):
    history = list(history)
    if len(history) > 10:
        history = history[-10:]
    question = []
    answers = []
    for i in history:
        question.append(i[0])
        answers.append(i[1])
    return (question, answers)


def delet_path(user):
    path = os.path.join(settings.STATIC_UPLOAD_DIR, user)
    if os.path.exists(path) and os.path.isdir(path):
        shutil.rmtree(path)
        print(f"Directory '{path}' deleted successfully.")
    else:
        print(f"Directory '{path}' does not exist.")


@csrf_exempt
@api_view(["DELETE"])
# @permission_classes([AllowAny])
def delet_file(request, user, filename, fileType):
    base_path = settings.STATIC_UPLOAD_DIR
    if fileType == "akt":
        base_path = settings.INTERNAL_ACT

    file_path = os.path.join(base_path, user, filename)
    csv_path = os.path.join(
        base_path, user, filename.rsplit(".", 1)[0] + ".csv")

    if os.path.exists(file_path) and os.path.exists(csv_path):
        os.remove(file_path)
        os.remove(csv_path)
        print(f"File '{file_path}' has been deleted.")
        return JsonResponse({"message": "File deleted successfully"})
    else:
        return JsonResponse({"error": "File not found"}, status=404)


def delet_photo(user):
    user_photos_path = os.path.join(
        settings.BASE_DIR, "static/userphotos", f"{user}.png")
    if os.path.exists(user_photos_path):
        os.remove(user_photos_path)
        print(f"File '{user_photos_path}' has been deleted.")
    else:
        print(f"File '{user_photos_path}' does not exist.")


def get_files(request, user, fileType):
    base_path = settings.STATIC_UPLOAD_DIR
    if fileType == "akt":
        base_path = settings.INTERNAL_ACT

    mypath = base_path + "/" + user
    pdf_files = allFileformat(mypath, ".pdf")
    docx_files = allFileformat(mypath, ".docx")

    all_files = pdf_files + docx_files
    return JsonResponse({"users": all_files})


def delet_user(request, user):
    find = UserValues.objects.filter(user=user)
    find.delete()
    userHistoy = History.objects.filter(sender=user)
    userHistoy.delete()

    users = getalldirs(settings.STATIC_UPLOAD_DIR)

    if user in users:
        index = users.index(user)
        wanted_users = users[index]
        delet_path(wanted_users)
        delet_photo(user)
    return redirect("/")


def manage_user(request, user):
    if request.method == "POST":
        form = UserValueForm(request.POST)
        if form.is_valid():  # Validate the form first
            find = UserValues.objects.filter(user=user)
            find.delete()
            chat_message = UserValues(
                user=user,
                splitter=form.cleaned_data["splitter"],
                chunksize=form.cleaned_data["chunksize"],
                overlap=form.cleaned_data["overlap"],
                temp=form.cleaned_data["temp"],
            )
            chat_message.save()
            reembedfiles(user)
        return redirect(f"/chat/{user}")
    else:
        form = UserValueForm()

    context = {"user": user, "form": form}
    return render(request, "manage-user.html", context)


def manage_users(request):
    mypath = settings.STATIC_UPLOAD_DIR
    upload_dir = os.path.join(settings.BASE_DIR, "static/userphotos")
    users = getalldirs(mypath)
    userphotos = allFileformat(upload_dir, ".png")
    userphotos.sort()
    users.sort()
    combined = zip(users, userphotos)
    context = {"combined": combined}

    return render(request, "manage-users.html", context)


def askingT(user, query):
    parser = Parsers(settings.OPENAI_KEY)
    query_vector = parser.embedquerry(query)

    chunks, vectors, contract_info = system_file_parser(query_vector, user)
    history = user_history(user)
    pastQuestion, pastAnswe = unpack_history(history)
    res = context_aware_responses(
        query, pastQuestion, pastAnswe, contract_info, user)
    return (res, contract_info)


def get_all_base_laws():
    files = []
    excluded_files = [
        "LIGJI__NR._03_L-212_I_PUNËS",
        "LIGJI__NR",
        "Draft Akti i Brendshëm - Kosove (7) (1)",
    ]
    mypath = settings.BASE_LAWS
    for dirpath, dirnames, filenames in walk(mypath):
        csv_files = [
            file
            for file in filenames
            if file.split(".")[0] not in excluded_files and file.endswith(".csv")
        ]
        files.extend(csv_files)
        break
    return files


def asking(user, query):
    parser = Parsers(settings.OPENAI_KEY)
    query_vector = parser.embedquerry(query)
    top_law_chunks, top_law_vector, sorted_law_files = system_file_parser_law(
        query_vector, user
    )  # Get_law_data(query_vector)
    #top_law_chunks2, top_law_vector2 = Get_law_data(query_vector)
    top_ak_chunks, top_ak_chnks, sorted_akt_files = system_file_parserT(
        query_vector, user, settings.INTERNAL_ACT
    )  # Get_aktin_i_brendshem_data(query_vector)

    chunks, vectors, contract_info = system_file_parser(top_law_vector, user)


    for i in contract_info:
        print(i)
        print(len(contract_info[i]["chunks"]))
        print(len(contract_info[i]["vectors"]))
    print("-" * 10)

    print(len(contract_info), "this is length 2<->")

    res, res_type = context_aware_responses(
        query=query,
        law_data=top_law_chunks,
        ak_data=top_ak_chunks,
        case_data=top_case_chunks,
        data=contract_info,
        user=user,
    )
    return (res, contract_info, res_type)


def user_history(user):
    chat_history = History.objects.filter(sender=user)
    question = []
    answers = []
    for i in chat_history:
        question.append(i.question)
        answers.append(i.respons)

    return zip(question, answers)


def getchunksforQuestin(request, user, question):

    instance = History.objects.filter(sender=user, question=question).first()

    data = []
    for i in instance.chunks.all():
        data.append(i.chunk_text)

    print("this is data and this is the len ", len(data))
    print(data)
    context = {
        #   'chunks': chunk_list,
        "data": data
    }
    return render(request, "test.html", context)


def saveHitoryChunsk(instance, data):

    print("$" * 50)
    for i in unpackdick(data):
        for j in i:
            print(j[:50])
            Chunk.objects.create(history=instance, chunk_text=j)
    print("=" * 50)
    for i in instance.chunks.all():
        print(i.chunk_text[:50])

    print("$" * 50)
    return None


def chat_leagl(request, user):
    res = []
    files = []
    combined = []
    responds = []
    context = {"user": user, "answer": responds,
               "files": files, "combined": combined}

    return render(request, "chat.html", context)


def getuser(username):
    return User.objects.get(username=username)


def get_context_data(text, user):

    contract_state = ContractState.objects.filter(user=getuser(user)).first()
    field_list = contract_state.fields  # ["Name", "Address", "Date of Birth"]

    # Format nicely for the system prompt
    fields_text = "\n".join(f"- {field}" for field in field_list)
    system_prompt = f"""
    Ti je një asistent ligjor. Nxirr të dhënat e mëposhtme nga mesazhi i përdoruesit dhe ktheji në JSON:

    {fields_text}

    Kthe vetëm JSON. Nëse ndonjë vlerë mungon, vendose si null.
    """

    client = OpenAI(api_key=settings.OPENAI_KEY)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"{text}"},
        ],
        temperature=0,
    )

    return response.choices[0].message.content


def genert_proerpt_contrext2(text):

    chunks = "\n".join(
        [f"\n{chunk}\n```" for chunk in settings.FILE["chunks"]])
    system_prompt = f"""
    You are a legal assistant that drafts professional contracts.
    Below you have a contract template (or parts of it):
    {chunks}

    Your task:
    - Use this template as the base.
    - Fill in the blank fields with the information provided by the user.
    - Ensure the contract is complete, clear, and written in a professional legal tone.
    - Clearly specify parties, dates, obligations, and any legal references.
    - Keep the contract’s structure (titles, articles, sub-articles, signatures).
    - If any required information is missing, leave a placeholder like "_____".
    
    Return only the finalized contract, ready for legal use.
    """

    print("99999" * 10)
    print(chunks)
    client = OpenAI(api_key=settings.OPENAI_KEY)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"{text}"},
        ],
        temperature=0,
    )

    return response.choices[0].message.content


def clean_and_load(content: str):
    # Remove triple backticks and optional language tag
    content = re.sub(r"^```(?:json)?\s*", "", content.strip())
    content = re.sub(r"\s*```$", "", content)
    return json.loads(content)


def json_to_docx(data: dict, filename="legal.docx", subfolder="my_docs"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    output_dir = os.path.join(settings.GENERATED_FILES, subfolder)
    os.makedirs(output_dir, exist_ok=True)
    full_path = os.path.join(output_dir, filename)

    print(settings.GENERATED_FILES)
    print(output_dir, " <=> ", full_path)
    print("json_to_docx -> {", data, "}")
    for section in data:
        if section["type"] == "heading":
            doc.add_heading(section["text"], level=section.get("level", 1))

        elif section["type"] == "paragraph":
            para = doc.add_paragraph(section["text"], style="Normal")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif section["type"] == "list":
            for item in section["items"]:
                doc.add_paragraph(item, style="List Bullet")

        elif section["type"] == "table":
            doc.add_heading(section.get("title", ""), level=2)
            table = doc.add_table(rows=0, cols=2)
            for row in section["rows"]:  # ✅ loop through the rows from JSON
                cells = table.add_row().cells
                cells[0].text = str(row.get("field", ""))
                cells[1].text = str(row.get("value", "") or "")

        elif section["type"] == "signature_block":
            table = doc.add_table(rows=1, cols=len(section["fields"]))
            for i, field in enumerate(section["fields"]):
                table.rows[0].cells[i].text = f"{field}: __________"

    doc.save(full_path)
    return full_path


def generate_legal_doc_json(text: str, user: str):
    print("JSON CALLED")

    contract_state = ContractState.objects.get(user=getuser(user))
    # <-- deserialize string back to dict
    file_data = json.loads(contract_state.file)
    chunks = "\n".join(file_data.get("chunks", []))
    print("$$$$" * 20)
    system_prompt = """
You are a legal document generator that outputs in strict JSON format. 
Your task is to take unstructured input and generate a **formal legal document** in Albanian (or English if specified) with a professional and legally accurate tone.

Rules:
1. Always detect the type of document requested (e.g., Ankesa, Kontratë Pune, Marrëveshje Bashkëpunimi, NDA, Kërkesë Zyrtare, Autorizim, etc.).
2. Use structured JSON objects for formatting:
   - {"type":"heading","level":1,"text":"..."} → For the main title.
   - {"type":"heading","level":2,"text":"..."} → For section titles.
   - {"type":"paragraph","text":"..."} → For normal paragraphs.
   - {"type":"list","items":["...","..."]} → For numbered or bullet lists.
   - {"type":"table","rows":[{"Field":"Value"}]} → For participant details (Palët, Emri, Pozita, Data, Email, Institucioni, etc.).
   - {"type":"signature_block","fields":["Data","Nënshkrimi"]} → For signatures.
3. Mandatory structure (adapt based on document type):
   - Title (e.g., "Kontratë Pune", "Marrëveshje për MosZbulim (NDA)", "Kërkesë për Rishikim", etc.)
   - Pjesëmarrësit / Palët (as a table with fields)
   - Data e Dokumentit (table row with date)
   - Përshkrimi ose Termat Kryesorë (paragraphs or lists depending on input)
   - Detyrimet dhe të Drejtat (paragraphs or lists, for contracts/agreements)
   - Dokumentacioni Shtesë i Bashkangjitur (list if provided)
   - Kërkesat ose Klauzolat (paragraphs/lists, depending on type)
   - Nënshkrimet (signature_block, with placeholders if missing)
4. Always keep a **formal legal tone** in all text.
5. If data is missing, insert placeholders like "________________".
6. The output must be a valid JSON array. Return **only JSON** — no commentary, no markdown, no explanations.

Examples of adaptability:
- If input = complaint → Output JSON representing a formal ankesa letter.
- If input = employment agreement → Output JSON for a "Kontratë Pune" with sections for palët, detyrimet, afati, pagesa, nënshkrimet.
- If input = NDA → Output JSON for a "Marrëveshje MosZbulimi" with clauses on confidentiality, exceptions, duration, governing law, signatures.

"""

    client = OpenAI(api_key=settings.OPENAI_KEY)
    response = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text},
        ],
    )

    # Parse GPT’s JSON output safely
    content = response.choices[0].message.content
    try:
        return clean_and_load(content)
        print("JSON YE ")

    except json.JSONDecodeError:
        print("JSON NO ")
        raise ValueError("GPT did not return valid JSON:\n" + content)


def genert_proerpt_contrext(text, user):
    contract_state = ContractState.objects.get(user=getuser(user))
    # <-- deserialize string back to dict
    file_data = json.loads(contract_state.file)
    chunks = "\n".join(file_data.get("chunks", []))

    system_prompt = f"""
You are a legal assistant that drafts professional contracts and complaints.

Below is a base template (or parts of it) for a legal complaint:
{chunks}

Your task:
- Use this template as a starting point, but feel free to expand or restructure sections if necessary to make the complaint complete and professional.
- Fill in the blank fields with the information provided by the user.
- Add details where appropriate, such as reasons for the complaint, suggested actions, or procedural notes.
- Ensure clarity, professional legal tone, and proper structure.
- Use inline fields or placeholders like "_____" for any missing information.
- Do NOT output bullet points unless absolutely necessary; format text naturally.
- Return only the final document, ready to export to DOCX.

"""

    client = OpenAI(api_key=settings.OPENAI_KEY)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text},
        ],
        temperature=0,
    )

    return response.choices[0].message.content


def write_string_to_pdf(content: str, filename: str = "demo.pdf", subfolder: str = "my_docs"):
    print("wtff" * 100)
    print("(##(#")
    output_dir = os.path.join(settings.BASE_DIR, subfolder)
    os.makedirs(output_dir, exist_ok=True)
    full_path = os.path.join(output_dir, filename)

    # PDF setup
    doc = SimpleDocTemplate(full_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    heading_style = ParagraphStyle(
        "Heading", parent=styles["Heading2"], spaceAfter=10)
    normal_style = styles["Normal"]
    bullet_style = styles["Bullet"]

    lines = [line.strip() for line in content.splitlines()
             if line.strip()]  # remove empty lines

    for line in lines:
        if line.startswith("- "):
            story.append(Paragraph(line[2:], bullet_style))
        elif line.startswith("**") and line.endswith("**"):
            story.append(Paragraph(line.strip("*"), heading_style))
        else:
            story.append(Paragraph(line, normal_style))
        story.append(Spacer(1, 6))  # small spacing

    # Build PDF
    doc.build(story)
    print(f"Saved to {full_path}")
    return full_path


def write_string_to_docx(content: str, filename: str = "demo.docx", subfolder: str = "my_docs"):
    output_dir = os.path.join(settings.BASE_DIR, subfolder)
    os.makedirs(output_dir, exist_ok=True)
    full_path = os.path.join(output_dir, filename)

    document = Document()
    lines = [line.strip() for line in content.splitlines()
             if line.strip()]  # remove empty lines

    for line in lines:
        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            # Bold headings
            document.add_paragraph(line.strip("**"), style="Heading 2")
        else:
            document.add_paragraph(line)

    document.save(full_path)
    print(f"Saved to {full_path}")
    return full_path


def review_one_file(request, user, fileName):
    print("hello123")

    client = OpenAI(api_key=settings.OPENAI_KEY)
    messages2 = [
        {
            "role": "system",
            "content": (
                "You are a highly skilled legal assistant specializing in legal document review.\n\n"
                "Your task is to review all provided legal documents carefully and return a clear, structured legal analysis. Focus specifically on:\n\n"
                "- Identifying potential legal risks or liabilities\n"
                "- Noting missing clauses or provisions\n"
                "- Highlighting compliance issues or inconsistencies\n"
                "- Pointing out ambiguous or vague language\n"
                "- Detecting contradictions across documents\n\n"
                "If the user hasn’t asked for a specific focus, perform a general legal review. "
                "If a review scope is provided (e.g., employment law, data privacy), prioritize findings related to that.\n\n"
                "Your tone should be professional, neutral, and clear. Use bullet points, sections, or numbered lists where appropriate.\n\n"
                "⚠️ Do **not** summarize the documents unless explicitly asked.\n"
                "⚠️ Do **not** provide general legal theory — focus on what's present or missing in the provided text.\n\n"
                "If something is unclear, flag it. If everything looks good, say so explicitly.\n\n"
                "You do not need to ask follow-up questions unless the prompt requests clarification.\n\n"
                "Example output structure:\n"
                "-------------------------\n"
                "Legal Review:\n\n"
                "1. Missing Clauses\n"
                "   - No termination clause detected.\n"
                "   - Lacks confidentiality agreement section.\n\n"
                "   - if somthing can be change specifke it and the valid change base on it \n"
                "2. Inconsistencies\n"
                "   - Section 4 conflicts with Section 2 regarding payment terms.\n\n"
                "3. Risk Assessment\n"
                "   - Clause 8 creates a potential liability due to vague language.\n\n"
                "   - specife if you can the Level of the risk Low/Medium/High"
                "4. General Observations\n"
                "   - Formatting is inconsistent.\n"
                "   - Passive voice used in critical obligations.\n\n"
                "End your review with a one-line summary of overall document quality "
                "(e.g., “The document is mostly complete but contains moderate risk areas.”)"
            ),
        },
    ]
    messages2.insert(
        1,
        {
            "role": "system",
            "content": (
                "The following guidance document outlines the standards and best practices for legal drafting and review in Kosovo. "
                "Refer to it strictly in all assessments:\n\n"
                # <- load this from the PDF or pre-parsed text
                f"{add_guidelines()}"
            ),
        },
    )

    print("LAW reveiw stared ⚠️ ->", user, "->", fileName)

    fileName = fileName[: fileName.rindex(".")]
    print(fileName)
    # addReviewContext(user, messages2)
    addReview_one(user, fileName, messages2)
    res = client.chat.completions.create(
        model="gpt-4o",
        messages=messages2,
        max_tokens=1500,  # test on the higer end to the lowest 50-600
        temperature=0.2,  # Strict and deterministic responses
    )

    response_message = res.choices[0].message

    return JsonResponse({"result": response_message.content})


def write_string_to_docx2(content: str, filename: str = "demo.docx", subfolder: str = "my_docs"):
    # Create full path: BASE_DIR / my_docs / filename
    output_dir = os.path.join(settings.BASE_DIR, subfolder)
    os.makedirs(output_dir, exist_ok=True)  # Ensure folder exists

    full_path = os.path.join(output_dir, filename)

    document = Document()
    for line in content.strip().splitlines():
        document.add_paragraph(line)
    document.save(full_path)

    print(f"Saved to {full_path}")
    return full_path


def download_file(request, user):
    print(settings.GENERATED_FILES)
    print(settings.STATIC_UPLOAD_DIR)
    print(settings.INTERNAL_ACT)
    file_path = os.path.join(settings.GENERATED_FILES, user, "legal.docx")
    return FileResponse(open(file_path, "rb"), as_attachment=True, filename="legal.docx")


# 🔥 Example usage:


@csrf_exempt  # remove if you handle CSRF with tokens
def chat_front(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST requests are allowed"}, status=405)

    try:
        data = json.loads(request.body)
        user = data.get("user")
        text = data.get("query")  # question from POST body
        text = unquote(text)

        print("TEXTTEXT")
        print(text)
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    if not user or not text:
        return JsonResponse({"error": "Missing 'user' or 'query' in request body"}, status=400)

    responds = ""
    genered = False
    review = False
    mypath = settings.STATIC_UPLOAD_DIR + "/" + user

    state = ContractState.objects.filter(user=getuser(user)).first()

    if not state or not state.user_info:
        responds, all_data, responds_type = asking(user, text)
        print("IMPORTANT ->", responds)
    else:
        user_contract_data = get_context_data(text, user)
        print("final boss !!!!" * 100)
        print("YAYYY WE EXTREAD ALL THE INFO WE NEED IT YAYYYY ", user_contract_data)

        responds, all_data, responds_type = asking(user, text)
        if responds_type != "Review":
            responds = genert_proerpt_contrext(user_contract_data, user)

            json_respons = generate_legal_doc_json(responds, user)
            savedocx = json_to_docx(
                json_respons, filename="legal.docx", subfolder=user)
            print(savedocx, "IS THIS THE RELEVT PATH ??")
            genered = True

        state.fields = ""
        state.file = ""
        state.user_info = False
        state.save()

    responds2 = ""
    if responds_type == "Review":
        responds2 = "\n-----\n".join(
            f"Batch {batch['batch']}:\nFiles: {', '.join(batch['files'])}\n\n{batch['review']}"
            for batch in responds
        )
    if responds_type == "Review":
        chat_message = History(
            sender=user,
            question=text,
            respons=responds2,
        )
        chat_message.save()
        saveHitoryChunsk(chat_message, all_data)
    else:
        chat_message = History(
            sender=user,
            question=text,
            respons=responds,
        )
        chat_message.save()
        saveHitoryChunsk(chat_message, all_data)

    pdf_files = allFileformat(mypath, ".pdf")
    word_files = allFileformat(mypath, ".docx")
    files = pdf_files + word_files

    combined = [{"question": q, "answer": a} for q, a in user_history(user)]

    if responds_type == "Review":

        print("2<=>2 " * 10)
        print(len(responds))
        print("2<=>2 " * 10)
    return JsonResponse(
        {"answer": responds, "history": combined,
            "Generted": genered, "type": responds_type}
    )


def chat_front2(request, user, query):
    responds = ""
    genered = False
    mypath = settings.STATIC_UPLOAD_DIR + "/" + user
    text = query  # question passed via URL path
    if text:
        state = ContractState.objects.filter(user=getuser(user)).first()

        if not state or not state.user_info:
            responds, all_data = asking(user, text)
        else:
            user_contract_data = get_context_data(text, user)
            print("final boss !!!!" * 100)
            print("YAYYY WE EXTREAD ALL THE INFO WE NEED IT YAYYYY ",
                  user_contract_data)

            responds, all_data = asking(user, text)
            responds = genert_proerpt_contrext(user_contract_data, user)

            json_respons = generate_legal_doc_json(responds, user)

            savedocx = json_to_docx(
                json_respons, filename="legal.docx", subfolder=user)
            print(savedocx, "IS THIS THE RELEVT PATH ??")

            # save = write_string_to_docx(responds)
            # savepdf = write_string_to_pdf(responds)
            # Save history
            state.fields = ""
            state.file = ""  # if you also want to clear the file content
            state.user_info = False
            state.save()
            genered = True

        chat_message = History(
            sender=user,
            question=text,
            respons=responds,
        )
        chat_message.save()
        saveHitoryChunsk(chat_message, all_data)

        pdf_files = allFileformat(mypath, ".pdf")
        word_files = allFileformat(mypath, ".docx")
        files = pdf_files + word_files

        combined = [{"question": q, "answer": a}
                    for q, a in user_history(user)]

    return JsonResponse({"answer": responds, "history": combined, "Generted": genered})


@csrf_exempt
def user_history_json(request, user):
    combined = [{"question": q, "answer": a} for q, a in user_history(user)]
    return JsonResponse({"user": user, "history": combined})


def chat(request, user):
    responds = ""
    mypath = settings.STATIC_UPLOAD_DIR + "/" + user
    if request.method == "POST":
        text = request.POST.get("question")
        state = ContractState.objects.filter(user=getuser(user)).first()
        if not state or not state.user_info:
            responds, all_data = asking(user, text)
        else:
            state.user_info = False
            state.save()
            state2 = ContractState.objects.filter(user=getuser(user)).first()
            user_contract_data = get_context_data(text, user)
            print("final boss !!!!" * 100)
            print("YAYYY WE EXTREAD ALL THE INFO WE NEED IT YAYYYY ",
                  user_contract_data)
            # responds, all_data = genert_template(user, user_contract_data):

            responds, all_data = asking(user, text)
            responds = genert_proerpt_contrext(user_contract_data, user)

            json_respons = generate_legal_doc_json(responds, user)

            savedocx = json_to_docx(json_respons)
            save = write_string_to_docx(responds)
            savepdf = write_string_to_pdf(responds)
            state.fields = ""
            state.file = ""  # if you also want to clear the file content
            state.user_info = False
            state.save()

        chat_message = History(
            # , chunks=unpackdick(contract_info))
            sender=user,
            question=text,
            respons=responds,
        )
        chat_message.save()
        saveHitoryChunsk(chat_message, all_data)
    pdf_files = allFileformat(mypath, ".pdf")
    word_files = allFileformat(mypath, ".docx")
    files = pdf_files + word_files

    combined = user_history(user)
    context = {"user": user, "answer": responds,
               "files": files, "combined": combined}
    return render(request, "chat.html", context)


"""
okay must make a comment formating patter for the chunks so it can be turend into data that i can parse very simple
so that i dont get a error then must change the chunk.chunks / data= json.loads(raw_response)
"""


def unpackdick(data):

    #    formatted_data = [
    #    {"file": filename, "chunks": info["chunks"]}

    #    for filename, info in data.items()
    # ]
    formatted_data = [info["chunks"] for filename, info in data.items()]

    return formatted_data


def home(request):
    print(get_all_base_laws())
    mypath = settings.STATIC_UPLOAD_DIR
    upload_dir = os.path.join(settings.BASE_DIR, "static/userphotos")
    users = getalldirs(mypath)
    userphotos = allFileformat(upload_dir, ".png")
    userphotos.sort()
    users.sort()
    combined = zip(users, userphotos)
    context = {"combined": combined}

    return render(request, "home.html", context)


def makedir_genert(user_name):
    # print(user_name)
    target_dir = os.path.join(settings.GENERATED_FILES, user_name)
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        print(f"Directory {user_name} created at {target_dir}")
    else:
        print(f"Directory {user_name} already exists at {target_dir}")
        os.makedirs(target_dir, exist_ok=True)


def makedir_internal_akt(user_name):
    # print(user_name)
    target_dir = os.path.join(settings.INTERNAL_ACT, user_name)
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        print(f"Directory {user_name} created at {target_dir}")
    else:
        print(f"Directory {user_name} already exists at {target_dir}")
        os.makedirs(target_dir, exist_ok=True)


def makedir(user_name):
    # print(user_name)
    target_dir = os.path.join(
        settings.BASE_DIR, "static", "uploads", user_name)
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        print(f"Directory {user_name} created at {target_dir}")
    else:
        print(f"Directory {user_name} already exists at {target_dir}")
        os.makedirs(target_dir, exist_ok=True)


def uploadphoto(photoname, photo):
    upload_dir = os.path.join(settings.BASE_DIR, "static/userphotos")

    # Save the file to the desired path
    fs = FileSystemStorage(location=upload_dir)
    fs.save(photoname, photo)
    answers = allFileformat(upload_dir, ".png")
    print(answers)


@csrf_exempt
def makedirForm(request):
    if request.method == "POST":
        form = MakeDirForm(request.POST, request.FILES)

        if form.is_valid():
            dirname = form.cleaned_data["name"]
            photo = form.cleaned_data["photo"]  # The uploaded image file

            makedir(dirname)
            print(dirname)

            photoname = dirname + photo.name[-4:]
            uploadphoto(photoname, photo)
            return redirect(f"/Manage-User/{dirname}")
    form = MakeDirForm()
    return render(request, "upload_file.html", {"form": form})


def save_file_akt(uploaded_file, user):
    fs = FileSystemStorage(location=settings.INTERNAL_ACT + f"/{user}")
    fs.save(uploaded_file.name, uploaded_file)

    file_url = f"{settings.INTERNAL_ACT}/{user}/{uploaded_file.name}"
    parser = Parsers(settings.OPENAI_KEY)
    user_value = UserValues.objects.filter(user=user).first()
    spliter = user_value.splitter
    chunksize = user_value.chunksize
    overlap = user_value.overlap
    # print(f'info about user {user} chunksize {
    #      chunksize} overlap {overlap} spliter {spliter}')
    parser.SetSpliter(spliter=spliter, chuncksize=chunksize, overlap=overlap)

    fileChunks, fileEmbedings = parser.embedd(file_url)

    parser.SaveCsv(
        settings.INTERNAL_ACT + "/" + user, uploaded_file.name, fileEmbedings, fileChunks
    )
    return fileEmbedings


def save_file(uploaded_file, user):
    fs = FileSystemStorage(location=settings.STATIC_UPLOAD_DIR + f"/{user}")
    fs.save(uploaded_file.name, uploaded_file)

    file_url = f"{settings.STATIC_UPLOAD_DIR}/{user}/{uploaded_file.name}"
    parser = Parsers(settings.OPENAI_KEY)
    user_value = UserValues.objects.filter(user=user).first()
    spliter = user_value.splitter
    chunksize = user_value.chunksize
    overlap = user_value.overlap
    # print(f'info about user {user} chunksize {
    #      chunksize} overlap {overlap} spliter {spliter}')
    parser.SetSpliter(spliter=spliter, chuncksize=chunksize, overlap=overlap)

    fileChunks, fileEmbedings = parser.embedd(file_url)

    parser.SaveCsv(
        settings.STATIC_UPLOAD_DIR + "/" +
        user, uploaded_file.name, fileEmbedings, fileChunks
    )
    return fileEmbedings


@csrf_exempt
def fileupload(request, user):
    if request.method == "POST" and request.FILES["file"]:
        uploaded_file = request.FILES["file"]
        my_file = Path(
            f"{settings.STATIC_UPLOAD_DIR}/{user}/{uploaded_file.name}")
        if not my_file.is_file():
            save_file(uploaded_file, user)
            return redirect(f"/chat/{user}/")
        else:
            return JsonResponse({"success": False, "data": "The File Existers allready"})
    else:
        form = FileUploadForm()
    return render(request, "save-static.html", {"form": form})
